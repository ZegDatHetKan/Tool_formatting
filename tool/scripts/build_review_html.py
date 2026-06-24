#!/usr/bin/env python3
"""Build the client review artifact from the REAL formatter output.

- fictional content -> LetterDocument -> render_letter -> out/lettera_fittizia.docx
- LibreOffice -> out/lettera_fittizia.pdf  (real, paginated: the deliverables)
- a CONTINUOUS copy (one tall page, no page breaks) -> PDF -> single PNG + bbox
  used ONLY for the web preview (letter shown as one long page)
- HTML "board": letter (permanent light per-section tints) on the left + section
  cards in a normal-flow column on the right + permanent light connector lines
  (SVG) joining each section to its card; hover/focus intensifies both. No hero.

Run:  .venv/bin/python tool/scripts/build_review_html.py
"""
from __future__ import annotations
import os, sys, html, glob, re, shutil, subprocess, tempfile
ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.insert(0, ROOT)
from docx import Document
from docx.shared import Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from formatters.letters import (LetterDocument, Span, render_letter,
                                paragraph, section_heading, subheading, list_item)

OUT = os.path.join(ROOT, "out"); OUTH = os.path.join(OUT, "html"); os.makedirs(OUTH, exist_ok=True)
TEMPLATE = os.path.join(ROOT, "assets", "Template_Vuoto.docx")
DOCX = os.path.join(OUT, "lettera_fittizia.docx")
PDF = os.path.join(OUT, "lettera_fittizia.pdf")
CONT_DOCX = os.path.join(OUTH, "_continuous.docx")
HTMLF = os.path.join(OUTH, "lettera_fittizia_review.html")
PNG = os.path.join(OUTH, "lettera_fittizia_review.png")
PT_CM = 28.3464567

# ---------- 1) fictional content ----------
DELIVERY = "A mezzo PEC, anticipata a mezzo posta elettronica"
DATE = "Bergamo, 24 giugno 2026"
RECIPIENT = ["Egr.", Span("Sig. Mario Rossi", bold=True), "Via dei Tigli, n. 14",
             "24121 Bergamo (BG)", Span("PEC: mario.rossi@esempio-pec.it", bold=True)]
SUBJECT = "OGGETTO: Diffida ad adempiere - contratto di fornitura del 10 gennaio 2026"
OPENING = "Egregio Signore,"
CLOSING = "In attesa di un cortese e sollecito riscontro, si porgono distinti saluti."
SIGNATURE = ["Bergamo Legal Società tra Avvocati S.r.l.",
             Span("Avv. Matteo Bertocchi", bold=True), "_______________________________"]
ATTACH = ["All. 1 - copia del contratto di fornitura del 10 gennaio 2026;",
          "All. 2 - copia della fattura n. 42."]
POST = "(documento fittizio, predisposto a solo scopo di revisione formale)"
BODY = [
    ("body",    paragraph("lo Studio scrivente agisce in nome e per conto di Aurora Servizi S.r.l., con sede in Bergamo, in forza di incarico ricevuto, in relazione al rapporto di fornitura di seguito descritto.")),
    ("subhead", subheading("Fatto")),
    ("body",    paragraph("1. In data 10 gennaio 2026 le parti sottoscrivevano un contratto di fornitura di materiale per ufficio, con pagamento a 30 giorni dalla consegna.")),
    ("body",    paragraph("2. La fornitura veniva regolarmente eseguita e fatturata (fattura n. 42), ma il relativo corrispettivo è rimasto integralmente impagato alla scadenza convenuta.")),
    ("ritual",  section_heading("DIFFIDA")),
    ("body",    paragraph("la S.V., come sopra individuata, a voler:")),
    ("list",    list_item("(i) provvedere al pagamento della somma di euro 3.500,00 entro e non oltre 15 giorni dalla ricezione della presente;")),
    ("list",    list_item("(ii) trasmettere allo Studio scrivente la contabile dell'avvenuto pagamento.")),
    ("body",    paragraph("In difetto di adempimento nel termine indicato, si procederà nelle competenti sedi giudiziarie per il recupero del credito, oltre interessi e spese.")),
]
letter = LetterDocument(delivery_method=DELIVERY, date_place=DATE, recipient_block=RECIPIENT,
    subject=SUBJECT, opening=OPENING, body_blocks=[b for _, b in BODY],
    closing=CLOSING, signature_block=SIGNATURE, attachments=ATTACH, postscript=POST)
res = render_letter(letter, TEMPLATE, DOCX)
print(f"[1] DOCX {DOCX} ({res.paragraph_count} par)")

# ---------- 2) classify rendered paragraphs ----------
def norm(s): return re.sub(r"\s+", " ", (s or "").replace("–","-").replace("—","-")).strip().lower()
def blk_text(b):
    c=b.content
    if isinstance(c,str): return c
    if isinstance(c,Span): return c.text
    if isinstance(c,list): return "".join(x.text if isinstance(x,Span) else str(x) for x in c)
    return str(c)
doc = Document(DOCX)
AL_NAME={AL.LEFT:"left",AL.RIGHT:"right",AL.CENTER:"center",AL.JUSTIFY:"justify",None:"left"}
AL_IT={"left":"sinistra","right":"destra","center":"centrato","justify":"giustificato"}
known={norm(DELIVERY):"delivery",norm(DATE):"date","oggetto:":"ogglabel",
       norm(SUBJECT.split(":",1)[1]):"oggbody",norm(OPENING):"opening",
       norm(CLOSING):"closing","allegati:":"attach",norm(POST):"post"}
for k,b in BODY: known[norm(blk_text(b))]=k
for ln in SIGNATURE: known[norm(ln.text if isinstance(ln,Span) else ln)]="signature"
for ln in ATTACH: known[norm(ln)]="attach"
def classify(p):
    t=norm(p.text)
    if t in known: return known[t]
    li=p.paragraph_format.left_indent
    if li is not None and abs(Emu(li).cm-8.5)<0.2: return "recipient"
    return "body"
def pinfo(p):
    sizes=[r.font.size.pt for r in p.runs if r.text and r.font.size]
    fonts=[r.font.name for r in p.runs if r.text and r.font.name]; pf=p.paragraph_format
    return dict(al=AL_NAME.get(p.alignment,"left"),size=max(sizes) if sizes else 12,
        font=fonts[0] if fonts else "Times New Roman",anybold=any(r.bold for r in p.runs if r.text),
        italic=any(r.italic for r in p.runs if r.text),
        li=round(Emu(pf.left_indent).cm,2) if pf.left_indent is not None else 0,
        sb=pf.space_before.pt if pf.space_before is not None else 0,
        sa=pf.space_after.pt if pf.space_after is not None else 0,kwn=bool(pf.keep_with_next))
body_seq=[(classify(p),norm(p.text)) for p in doc.paragraphs if p.text.strip()]
agg={}
for p in doc.paragraphs:
    if p.text.strip(): agg.setdefault(classify(p),pinfo(p))
sec=doc.sections[0]
def part_lines(part):
    out=[]
    for p in part.paragraphs:
        if p.text.strip():
            szs=[r.font.size.pt for r in p.runs if r.text and r.font.size]
            out.append((p.text,max(szs) if szs else None,p.runs[0].font.name if p.runs and p.runs[0].font.name else None))
    return out
header_lines, footer_lines = part_lines(sec.header), part_lines(sec.footer)
top_cm=round(Emu(sec.top_margin).cm,2); bot_cm=round(Emu(sec.bottom_margin).cm,2)
hdr_set=[norm(t) for t,_,_ in header_lines]
top_pt=top_cm*PT_CM; bot_pt=bot_cm*PT_CM

# ---------- helpers ----------
def soffice_pdf(src,outdir):
    prof=tempfile.mkdtemp(prefix="lo_")
    subprocess.run(["soffice","--headless",f"-env:UserInstallation=file://{prof}",
        "--convert-to","pdf","--outdir",outdir,src],check=True,
        stdout=subprocess.DEVNULL,stderr=subprocess.DEVNULL,timeout=180)
    shutil.rmtree(prof,ignore_errors=True)
    return os.path.join(outdir,os.path.splitext(os.path.basename(src))[0]+".pdf")
def bbox_pages(pdf):
    tmp=pdf+".bbox.xhtml"
    subprocess.run(["pdftotext","-bbox",pdf,tmp],check=True,timeout=60)
    x=open(tmp,encoding="utf-8").read(); os.remove(tmp)
    out=[]
    for pw,ph,body in re.findall(r'<page width="([\d.]+)" height="([\d.]+)">(.*?)</page>', x, re.S):
        words=[(float(a),float(b),float(c),float(d),html.unescape(e))
               for a,b,c,d,e in re.findall(r'<word xMin="([\d.]+)" yMin="([\d.]+)" xMax="([\d.]+)" yMax="([\d.]+)">([^<]*)</word>', body)]
        out.append((float(pw),float(ph),words))
    return out

soffice_pdf(DOCX, OUT)
print(f"[2] PDF {PDF}")
pag = bbox_pages(PDF)

# ---------- 3) continuous one-page copy ----------
H_body=0.0
for pw,ph,words in pag:
    body_w=[w for w in words if (top_pt-8) < w[1] < (ph-bot_pt-8)]
    if body_w: H_body += max(w[3] for w in body_w) - top_pt
new_h_pt = top_pt + H_body + bot_pt + 40
cont_pdf=None; npages=99
for _ in range(8):
    cont = Document(DOCX); cont.sections[0].page_height = Cm(new_h_pt/PT_CM); cont.save(CONT_DOCX)
    cont_pdf = soffice_pdf(CONT_DOCX, OUTH); npages=len(bbox_pages(cont_pdf))
    if npages<=1: break
    new_h_pt += 3*PT_CM
print(f"[3] continuous page {new_h_pt/PT_CM:.1f} cm -> {npages} pag")

# ---------- 4) single PNG + coords ----------
for old in glob.glob(os.path.join(OUTH,"lettera_fittizia_review-*.png")): os.remove(old)
if os.path.exists(PNG): os.remove(PNG)
subprocess.run(["pdftoppm","-png","-r","150","-singlefile",cont_pdf,os.path.splitext(PNG)[0]],check=True,timeout=120)
cpw,cph,cwords = bbox_pages(cont_pdf)[0]
os.remove(cont_pdf); os.remove(CONT_DOCX)

# ---------- 5) words -> lines -> regions ----------
def line_key(text,yMin,yMax,pageH):
    if yMax <= top_pt+8:
        for t in hdr_set:
            if t and (t in text or text in t): return "header"
        return "header"
    if yMin >= pageH-bot_pt-8: return "footer"
    for key,ptext in body_seq:
        if text and text==ptext: return key
    for key,ptext in body_seq:
        if len(text)>=4 and text in ptext: return key
    return "body"
cwords.sort(key=lambda w:(round(w[1],1),w[0]))
lines=[]; cur=[]
for w in cwords:
    if cur and abs(w[1]-cur[-1][1])>5: lines.append(cur); cur=[]
    cur.append(w)
if cur: lines.append(cur)
L=[]
for ln in lines:
    xs=[w[0] for w in ln]+[w[2] for w in ln]; ys=[w[1] for w in ln]+[w[3] for w in ln]
    txt=norm(" ".join(w[4] for w in sorted(ln,key=lambda w:w[0])))
    L.append((min(xs),min(ys),max(xs),max(ys),txt))
regions=[]; run=None
for (x0,y0,x1,y1,txt) in L:
    k=line_key(txt,y0,y1,cph); lh=y1-y0
    if run and run["k"]==k and (y0-run["y1"])<=lh*1.8:
        run["x0"]=min(run["x0"],x0);run["y0"]=min(run["y0"],y0);run["x1"]=max(run["x1"],x1);run["y1"]=max(run["y1"],y1)
    else:
        if run: regions.append(run)
        run=dict(k=k,x0=x0,y0=y0,x1=x1,y1=y1)
if run: regions.append(run)
# merge regions of the same key into ONE card per key (keep all boxes for tints)
def pct(r,pad=1.2):
    return (max(0,(r["x0"]-pad)/cpw*100),max(0,(r["y0"]-pad)/cph*100),
            (r["x1"]-r["x0"]+2*pad)/cpw*100,(r["y1"]-r["y0"]+2*pad)/cph*100)

# ---------- 6) metadata ----------
COLORS={"header":(100,116,139),"delivery":(202,138,4),"date":(13,148,136),
    "recipient":(79,70,229),"ogglabel":(225,29,72),"oggbody":(219,39,119),
    "opening":(22,163,74),"body":(37,99,235),"subhead":(124,58,237),"ritual":(234,88,12),
    "list":(8,145,178),"closing":(101,163,13),"signature":(146,64,14),"attach":(162,28,175),
    "post":(71,85,105),"footer":(100,116,139)}
LABEL={"header":"Intestazione (letterhead)","delivery":"Metodo di trasmissione","date":"Data e luogo",
 "recipient":"Destinatario","ogglabel":"OGGETTO — etichetta","oggbody":"OGGETTO — contenuto",
 "opening":"Apertura (saluto)","body":"Paragrafo di corpo","subhead":"Titoletto a sinistra",
 "ritual":"Titolo rituale centrato","list":"Voce di elenco","closing":"Congedo",
 "signature":"Blocco firma","attach":"Allegati","post":"Postilla / disclaimer","footer":"Piè di pagina"}
TAG={"header":"Ereditato","footer":"Ereditato","ogglabel":"Punto sensibile","oggbody":"Punto sensibile","signature":"Da approvare"}
NOTE={"header":"Carta intestata ereditata dal template; il corpo non la duplica.",
 "delivery":"Reso solo se fornito. Può anche essere la prima riga del destinatario.",
 "date":"Resa solo se fornita; lo script non inventa una data assente.",
 "recipient":"Appellativo personale unito al nome (“Egr. Sig. Mario Rossi”); “Spett.le” davanti a società resta su riga propria.",
 "ogglabel":"Punto sensibile (feedback 027): etichetta come titolo centrato, lasciata come scritta.",
 "oggbody":"La contestazione su 027 era il contenuto centrato invece che giustificato: ora è giustificato.",
 "opening":"Formale → grassetto + giustificato; e-mail → tondo + sinistra. Reso solo se fornito.",
 "body":"Normalizzazione: “–/—” → “-”; divisori “***” rimossi.",
 "subhead":"Es. “Fatto”, “MOTIVA …”; i verbi dispositivi usano titoletto + contenuto a capo.",
 "ritual":"Es. “IN FATTO”, “DICHIARA”, “DIFFIDA”, “PREMESSO CHE”, “INVITA”.",
 "list":"Marcatore ((i)/(a)/•) incluso nel testo.","closing":"Stesso lato della firma (destra).",
 "signature":"Nomi firmatari in grassetto; “(Firma digitale)” in corsivo; spaziatura ariosa — da approvare.",
 "attach":"“Allegati:” in grassetto; voci rientrate.","post":"Piccola, in corsivo.",
 "footer":"Dati di registrazione dello Studio, ereditati dal template."}
ORDER=["header","delivery","date","recipient","ogglabel","oggbody","opening","body",
 "subhead","ritual","list","closing","signature","attach","post","footer"]
def esc(s): return html.escape(s or "",quote=True)
def details(key):
    rows=[]
    if key in ("header","footer"):
        lines=header_lines if key=="header" else footer_lines
        szs=sorted({s for _,s,_ in lines if s}); font=(lines[0][2] if lines else "Garamond") or "Garamond"
        rows=[("Provenienza","Template (non generato)"),("Font",esc(font))]
        if szs: rows.append(("Dimensioni"," / ".join(f"{s:g} pt" for s in szs)))
        rows+=[("Allineamento","centrato"),("Pagina",f"margine sup. {top_cm:g} cm" if key=="header" else f"margine inf. {bot_cm:g} cm")]
    else:
        i=agg.get(key)
        if i:
            rows=[("Provenienza","Generato dallo script"),("Font",esc(i["font"])),("Dimensione",f"{i['size']:g} pt"),("Allineamento",AL_IT[i["al"]])]
            if i["li"]: rows.append(("Rientro sx",f"{i['li']:g} cm"))
            rows+=[("Grassetto","sì" if i["anybold"] else "no"),("Corsivo","sì" if i["italic"] else "no"),
                   ("Spaziatura",f"prima {i['sb']:g} / dopo {i['sa']:g} pt"),("Anti-orfano","keep-with-next" if i["kwn"] else "no")]
    dl="".join(f"<dt>{k}</dt><dd>{v}</dd>" for k,v in rows)
    return f"<dl>{dl}</dl>"+(f'<div class="hint">{NOTE.get(key,"")}</div>' if NOTE.get(key) else "")
def summary(key):
    if key in ("header","footer"): return "Ereditato dal template (Garamond)."
    i=agg.get(key)
    if not i: return ""
    bits=[AL_IT[i["al"]],f"{i['size']:g} pt"]
    if i["anybold"]: bits.append("grassetto")
    if i["italic"]: bits.append("corsivo")
    if i["li"]: bits.append(f"rientro {i['li']:g} cm")
    return ", ".join(bits)+"."

def col(k): r,g,b=COLORS[k]; return f"{r},{g},{b}"
# region tint divs (all boxes); the FIRST box per key gets an anchor id for the wire
seen=set(); rg_html=[]
for idx,r in enumerate(regions):
    k=r["k"]; l,t,w,h=pct(r)
    anchor = (k not in seen)
    aid = f' id="anc-{k}"' if anchor else ""
    seen.add(k)
    rg_html.append(f'<div class="rg{ " anchor" if anchor else ""}"{aid} data-k="{k}" tabindex="0" '
                   f'title="{esc(LABEL[k])}" style="left:{l:.2f}%;top:{t:.2f}%;width:{w:.2f}%;height:{h:.2f}%;--c:{col(k)}"></div>')
# one card per key, in document order of first appearance
order_seen=[];
for r in regions:
    if r["k"] not in order_seen: order_seen.append(r["k"])
# sections rendered ONLY if the author supplies them (otherwise omitted)
OPTIONAL={"delivery","date","opening","closing","attach","post"}
notes_html=[]
for k in order_seen:
    tag=f'<span class="tag">{TAG[k]}</span>' if k in TAG else ""
    opt=' <span class="opt" title="Opzionale: resa solo se fornita, altrimenti omessa">opzionale</span>' if k in OPTIONAL else ""
    notes_html.append(f'''      <div class="note" data-k="{k}" tabindex="0" style="--c:{col(k)}">
        <div class="nhd"><span class="dot"></span><span class="ntitle">{LABEL[k]}</span>{opt}{tag}</div>
        <div class="nsum">{summary(k)}</div>
        <div class="ndet">{details(k)}</div>
      </div>''')
print("[6] regioni:",len(regions)," sezioni:",order_seen)

PAGE=f"""<!DOCTYPE html>
<html lang="it"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1">
<title>Revisione formattazione — Bergamo Legal (lettera fittizia)</title>
<style>
  :root{{--ink:#1b2330;--muted:#65707f;--line:#e3e8ef;--bg:#eceff3}}
  *{{box-sizing:border-box}} html,body{{margin:0}}
  body{{font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,Helvetica,Arial,sans-serif;color:var(--ink);background:var(--bg);line-height:1.45}}
  .strip{{padding:10px 20px;border-bottom:1px solid var(--line);background:#fff;font-size:12px;color:var(--muted)}}
  .strip b{{color:var(--ink);font-weight:600}}
  .stage{{padding:28px 24px 80px;display:flex;justify-content:center}}
  .board{{position:relative;display:flex;align-items:flex-start;gap:96px}}
  #wires{{position:absolute;left:0;top:0;pointer-events:none;z-index:3;overflow:visible}}
  #wires path{{fill:none;stroke-width:1.3;opacity:.28;transition:opacity .12s,stroke-width .12s}}
  #wires path.on{{opacity:1;stroke-width:2.4}}
  .letter{{position:relative;flex:0 0 600px;width:600px;max-width:100%;line-height:0}}
  .letter img{{width:100%;display:block;border:1px solid var(--line);box-shadow:0 4px 22px rgba(20,30,45,.12);background:#fff}}
  .rg{{position:absolute;border-radius:3px;cursor:pointer;outline:none;background:rgba(var(--c),.11);transition:background .12s,box-shadow .12s}}
  .rg:hover,.rg:focus,.rg.on{{background:rgba(var(--c),.30);box-shadow:0 0 0 2px rgba(var(--c),.9)}}
  .notes{{flex:0 0 340px;width:340px;max-width:100%;display:flex;flex-direction:column;gap:7px;z-index:4}}
  .note{{background:#fff;border:1px solid var(--line);border-left:4px solid rgb(var(--c));border-radius:8px;
         padding:8px 11px;box-shadow:0 1px 3px rgba(20,30,45,.05);cursor:pointer;outline:none;
         transition:box-shadow .12s,background .12s,transform .12s}}
  .note:hover,.note:focus,.note.on{{background:rgba(var(--c),.07);box-shadow:0 4px 16px rgba(var(--c),.30);transform:translateX(-2px)}}
  .nhd{{display:flex;align-items:center;gap:7px}}
  .dot{{width:11px;height:11px;border-radius:3px;background:rgb(var(--c));flex:0 0 auto}}
  .ntitle{{font-weight:600;font-size:12.5px}}
  .tag{{font-size:9px;text-transform:uppercase;letter-spacing:.04em;font-weight:700;color:rgb(var(--c));margin-left:auto}}
  .opt{{font-size:9px;text-transform:uppercase;letter-spacing:.03em;color:var(--muted);border:1px dashed #b3bcca;border-radius:999px;padding:0 6px;white-space:nowrap}}
  .nsum{{font-size:11px;color:var(--muted);margin-top:3px}}
  .ndet{{max-height:0;overflow:hidden;transition:max-height .2s ease;font-size:11px}}
  .note:hover .ndet,.note:focus .ndet,.note.on .ndet{{max-height:340px;margin-top:7px;border-top:1px dashed rgba(var(--c),.5);padding-top:7px}}
  .ndet dl{{display:grid;grid-template-columns:auto 1fr;gap:1px 10px;margin:0}}
  .ndet dt{{color:var(--muted)}} .ndet dd{{margin:0}} .ndet .hint{{margin-top:5px;color:#475569}}
  @media (max-width:1040px){{
    .board{{flex-direction:column;gap:14px}} #wires{{display:none}}
    .letter{{flex:none;width:100%}} .notes{{flex:none;width:100%}}
    .ndet{{max-height:none;margin-top:7px;border-top:1px dashed rgba(var(--c),.5);padding-top:7px}}
  }}
  @media print{{body{{background:#fff}} #wires path{{opacity:.5}} .ndet{{max-height:none!important;margin-top:7px;border-top:1px dashed rgba(var(--c),.5);padding-top:7px}} .note{{box-shadow:none}}}}
</style></head>
<body>
<div class="strip">Lettera <b>fittizia</b> resa dal formatter (DOCX → PDF → immagine): font e impaginazione reali, <b>pagina unica</b>. <b>Passa il mouse</b> per un'anteprima, <b>clicca</b> una sezione o una scheda per <b>fissarla</b> (più sezioni restano aperte; riclicca per chiudere). Il segno <span class="opt">opzionale</span> indica i blocchi resi solo se forniti. Documenti puliti: <b>out/lettera_fittizia.docx</b> / <b>.pdf</b>.</div>
<div class="stage">
  <div class="board" id="board">
    <svg id="wires"></svg>
    <div class="letter" id="letter">
      <img id="page" src="{esc(os.path.basename(PNG))}" alt="Lettera fittizia (render reale, pagina unica)">
{os.linesep.join("      "+x for x in rg_html)}
    </div>
    <div class="notes" id="notes">
{os.linesep.join(notes_html)}
    </div>
  </div>
</div>
<script>
(function(){{
  function all(s){{return Array.prototype.slice.call(document.querySelectorAll(s));}}
  var board=document.getElementById('board'), wires=document.getElementById('wires'), img=document.getElementById('page');
  var pinned={{}}, hovered=null;                 // pinned = click-fixed; hovered = transient
  function anchor(k){{return document.getElementById('anc-'+k);}}
  function isOn(k){{return !!pinned[k] || hovered===k;}}
  function apply(){{                              // reflect state on regions, cards and lines
    all('.note').forEach(function(n){{n.classList.toggle('on',isOn(n.getAttribute('data-k')));}});
    all('.rg').forEach(function(e){{e.classList.toggle('on',isOn(e.getAttribute('data-k')));}});
    all('#wires path').forEach(function(p){{p.classList.toggle('on',isOn(p.getAttribute('data-k')));}});
  }}
  function draw(){{
    if(window.innerWidth<=1040){{ wires.innerHTML=''; return; }}
    var br=board.getBoundingClientRect();
    wires.setAttribute('width',board.clientWidth); wires.setAttribute('height',board.clientHeight);
    var s='';
    all('.note').forEach(function(n){{
      var k=n.getAttribute('data-k'), rg=anchor(k); if(!rg) return;
      var a=rg.getBoundingClientRect(), b=n.getBoundingClientRect();
      var x1=a.right-br.left, y1=a.top-br.top+a.height/2;
      var x2=b.left-br.left,  y2=b.top-br.top+b.height/2;
      var mx=(x1+x2)/2, c=getComputedStyle(n).getPropertyValue('--c').trim();
      s+='<path data-k="'+k+'" d="M'+x1+','+y1+' C'+mx+','+y1+' '+mx+','+y2+' '+x2+','+y2+'" stroke="rgb('+c+')"></path>';
    }});
    wires.innerHTML=s; apply();                   // KEY FIX: re-apply highlight after rebuild
  }}
  function later(){{ setTimeout(draw,220); }}      // after a card-expand transition
  function wire(el){{var k=el.getAttribute('data-k');
    el.addEventListener('mouseenter',function(){{hovered=k; apply(); draw(); later();}});
    el.addEventListener('mouseleave',function(){{hovered=null; apply(); later();}});
    el.addEventListener('focus',function(){{hovered=k; apply(); draw(); later();}});
    el.addEventListener('blur',function(){{hovered=null; apply(); later();}});
    el.addEventListener('click',function(e){{e.stopPropagation(); pinned[k]=!pinned[k]; apply(); draw(); later();}});
    el.addEventListener('keydown',function(e){{if(e.key==='Enter'||e.key===' '){{e.preventDefault(); pinned[k]=!pinned[k]; apply(); draw(); later();}}}});
  }}
  all('.rg').forEach(wire); all('.note').forEach(wire);
  function redraw(){{requestAnimationFrame(draw);}}
  if(img.complete) redraw(); else img.addEventListener('load',redraw);
  window.addEventListener('load',redraw);
  var t=null; window.addEventListener('resize',function(){{if(t)return;t=setTimeout(function(){{t=null;draw();}},120);}});
}})();
</script>
</body></html>
"""
open(HTMLF,"w",encoding="utf-8").write(PAGE)
print(f"[7] HTML {HTMLF}")
