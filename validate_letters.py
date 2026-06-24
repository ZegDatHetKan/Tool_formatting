#!/usr/bin/env python3
"""Validation harness for the Bergamo Legal letters formatter.

Run:  .venv/bin/python validate_letters.py

It demonstrates, end to end, that:
  1. the template opens;
  2. output DOCX files are created (under out/, never in previous_works/);
  3. the historical corpus in previous_works/ is left byte-for-byte untouched;
  4. at least three representative letter structures can be expressed through
     the semantic schema and rendered;
  5. each rendered output inherits the template letterhead/header/footer/margins
     and obeys the core typographic rules (Times New Roman on every run, no
     en/em dash, expected block alignments).

Exit code is non-zero if any check fails.
"""

from __future__ import annotations

import hashlib
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Emu

from formatters.letters import (
    LetterDocument,
    Span,
    SubjectStyle,
    disposition,
    list_item,
    paragraph,
    render_letter,
    section_heading,
    subheading,
)

ROOT = os.path.dirname(os.path.abspath(__file__))
TEMPLATE = os.path.join(ROOT, "assets", "Template_Vuoto.docx")
CORPUS = os.path.join(ROOT, "previous_works")
OUT = os.path.join(ROOT, "out")


# ---------------------------------------------------------------------------
# Representative letters (each exercises different parts of the schema)
# ---------------------------------------------------------------------------

def example_diffida() -> LetterDocument:
    """Formal diffida with Fatto/Diritto, ritual heading, lists, two signers
    (structure inspired by corpus 028/007)."""
    return LetterDocument(
        delivery_method="A mezzo Raccomandata A/R, anticipata a mezzo posta elettronica",
        date_place="Bergamo, 24 giugno 2026",
        recipient_block=[
            "Egr.",
            Span("Sig.ra Mario Esempio", bold=True),
            "Via di Prova, n. 1",
            "24100 Bergamo (BG)",
            Span("e-mail: mario.esempio@example.com", bold=True),
        ],
        subject="OGGETTO: Diffida alla rimozione di contenuti diffamatori",
        opening="Egregia Signora,",
        body_blocks=[
            paragraph("lo Studio scrivente agisce in nome e per conto di ESEMPIO S.r.l."),
            subheading("Fatto"),
            paragraph("1. Risulta al nostro Assistito che la S.V. ha pubblicato contenuti lesivi."),
            list_item("(i) primo contenuto contestato;"),
            list_item("(ii) secondo contenuto contestato;"),
            subheading("Diritto"),
            paragraph("2. Le affermazioni integrano gli estremi della diffamazione."),
            paragraph("Tutto cio premesso e considerato, lo Studio scrivente"),
            section_heading("DIFFIDA"),
            paragraph("la S.V., come sopra individuata, a voler:"),
            list_item("(i) rimuovere integralmente i contenuti entro 15 giorni;"),
            list_item("(ii) astenersi da ulteriori pubblicazioni."),
        ],
        closing="In attesa di un cortese, sollecito riscontro, si porgono distinti saluti.",
        signature_block=[
            Span("Avv. Matteo Bertocchi", bold=True),
            Span("Avv. Daiana Chiappa", bold=True),
        ],
    )


def example_email_comunicazione() -> LetterDocument:
    """E-mail/PEC-style communication (corpus 005/011/032): NO postal recipient
    block and NO date — only what the author sent is rendered (D2). Plain
    opening, inline subject, list, single signer."""
    return LetterDocument(
        # no recipient_block, no date_place: nothing is invented/completed
        subject="Oggetto: Re: Aggiornamento numeri D-U-N-S",
        opening="Gentile Stefania,",
        body_blocks=[
            paragraph("la ringrazio per il Suo riscontro e per il supporto."),
            paragraph("Le due entita coinvolte sono:"),
            list_item("Avvocato Matteo Bertocchi - D-U-N-S 440487417;"),
            list_item("Bergamo Legal Societa tra Avvocati s.r.l. - D-U-N-S 442296279;"),
            paragraph("Resto a disposizione per ulteriori chiarimenti."),
        ],
        closing="Cordiali saluti,",
        signature_block=[
            Span("Avv. Matteo Bertocchi", bold=True),
            "Amministratore Unico e socio di maggioranza",
            "Bergamo Legal Societa tra Avvocati s.r.l.",
        ],
        subject_style=SubjectStyle.INLINE,
        subject_content_center=True,
        opening_formal=False,
    )


def example_contestazione_pec() -> LetterDocument:
    """PEC contestation: delivery note top-right italic, date above recipient,
    split subject, plain firm + bold signer (inspired by corpus 003/029)."""
    return LetterDocument(
        delivery_method="Trasmessa a mezzo PEC",
        date_place="Bergamo, li 24 giugno 2026",
        recipient_block=[
            Span("Spett.le", bold=True),
            Span("Papi Solutions S.r.l.", bold=True),
            "Via Europa, 41 - 24035 Curno (BG)",
            Span("PEC: papisolutions@legalmail.it", bold=True),
        ],
        subject="OGGETTO: contestazione fattura n. 77 del 25.03.2026",
        opening="Spett.le Societa,",
        body_blocks=[
            paragraph("lo scrivente Studio assiste Service Italia S.a.s. (la Cliente)."),
            paragraph("La pretesa e contestata."),
            paragraph("La proposta e da intendersi quale offerta finale e non rinnovabile."),
        ],
        closing="In attesa di cortese e sollecito riscontro, si porgono distinti saluti.",
        signature_block=[
            "Bergamo Legal Societa tra Avvocati S.r.l.",
            Span("Avv. Matteo Bertocchi", bold=True),
        ],
        opening_formal=False,
    )


def example_informativa_with_placeholder() -> LetterDocument:
    """Informativa with date BELOW the recipient and a trailing postscript;
    deliberately keeps a placeholder to exercise needs_review (corpus 001)."""
    return LetterDocument(
        recipient_block=[
            Span("Egr./Gent.ma", italic=True),
            "Sig./Sig.ra [DA INSERIRE: NOME E COGNOME]",
            "Consigliere/a del Comune di Madone",
        ],
        date_place="Madone, li [DA INSERIRE: data]",
        subject="OGGETTO: informativa in ordine a fatti occorsi",
        opening="Egregio/a Consigliere/a,",
        body_blocks=[
            paragraph("con la presente il sottoscritto si rivolge a Lei."),
            section_heading("IN FATTO"),
            paragraph("1. Nel corso della seduta del Consiglio Comunale..."),
            section_heading("INVITA LA S.V."),
            paragraph("a voler assumere una formale presa di posizione."),
        ],
        closing="Con osservanza,",
        signature_block=[
            Span("Ing. Sergio Esempio", bold=True),
            "__________________________________",
        ],
        postscript="(trasmessa per il tramite dei difensori di fiducia - Bergamo Legal S.t.A.)",
        date_above_recipient=False,
    )


def example_istanza_ritiro() -> LetterDocument:
    """Formal istanza with NO salutation, a ritual centered verb (DICHIARA) and
    two dispositive 'titoletti' (MOTIVA, CHIEDE ALTRESÌ) whose content goes a
    capo (structure inspired by corpus 004). Fictional data."""
    return LetterDocument(
        recipient_block=[
            "Tramite servizio online / PEC a:",
            "ufficio.esempio@pec.example.it",
            "Spett.le",
            "Ufficio di Esempio",
        ],
        subject="OGGETTO: Ritiro della domanda di esempio n. 0000",
        opening=None,  # istanza formale: nessun saluto
        body_blocks=[
            paragraph("Esempio S.r.l., in persona del legale rappresentante,"),
            section_heading("DICHIARA"),
            paragraph("di voler ritirare la domanda di cui in oggetto."),
            subheading("MOTIVA il ritiro come segue."),
            paragraph("La richiedente ha rivisto la propria strategia."),
            *disposition(
                "CHIEDE ALTRESÌ",
                "il rimborso delle tasse versate, ove rimborsabili.",
            ),
        ],
        closing="In fede,",
        signature_block=[
            Span("Avv. Matteo Bertocchi", bold=True),
            "(Firma digitale)",
        ],
    )


EXAMPLES = {
    "diffida": example_diffida,
    "email_comunicazione": example_email_comunicazione,
    "contestazione_pec": example_contestazione_pec,
    "informativa": example_informativa_with_placeholder,
    "istanza_ritiro": example_istanza_ritiro,
}


# ---------------------------------------------------------------------------
# Checks
# ---------------------------------------------------------------------------

def corpus_fingerprint() -> dict[str, str]:
    fp = {}
    for name in sorted(os.listdir(CORPUS)):
        path = os.path.join(CORPUS, name)
        if os.path.isfile(path):
            with open(path, "rb") as fh:
                fp[name] = hashlib.sha256(fh.read()).hexdigest()
    return fp


def check_output(path: str) -> list[str]:
    """Re-open a rendered output and assert template inheritance + type rules."""
    errors: list[str] = []
    doc = Document(path)

    # 1. Letterhead preserved in the header.
    header_text = "\n".join(p.text for p in doc.sections[0].header.paragraphs)
    if "BERGAMO LEGAL" not in header_text:
        errors.append(f"{path}: letterhead 'BERGAMO LEGAL' assente nell'header")

    # 2. Footer (firm registration) preserved.
    footer_text = "\n".join(p.text for p in doc.sections[0].footer.paragraphs)
    if "P.IVA" not in footer_text and "R.E.A" not in footer_text:
        errors.append(f"{path}: footer dello Studio non preservato")

    # 3. Margins preserved (top ~5.91 cm).
    top_cm = Emu(doc.sections[0].top_margin).cm
    if not (5.8 < top_cm < 6.0):
        errors.append(f"{path}: margine superiore inatteso ({top_cm:.2f} cm)")

    # 4. Body actually composed.
    if len(doc.paragraphs) < 5:
        errors.append(f"{path}: corpo troppo corto ({len(doc.paragraphs)} paragrafi)")

    # 5. Times New Roman on every body run + no en/em dash.
    for p in doc.paragraphs:
        for r in p.runs:
            if r.text and r.font.name != "Times New Roman":
                errors.append(f"{path}: run non Times New Roman: {r.text[:30]!r}")
                break
        if "–" in p.text or "—" in p.text:
            errors.append(f"{path}: en/em dash residuo: {p.text[:40]!r}")

    # 6. A right-aligned signature paragraph exists.
    if not any(p.alignment == WD_ALIGN_PARAGRAPH.RIGHT and p.text.strip()
               for p in doc.paragraphs):
        errors.append(f"{path}: nessun paragrafo allineato a destra (firma/congedo)")

    return errors


def main() -> int:
    errors: list[str] = []

    # Template opens.
    try:
        Document(TEMPLATE)
        print(f"[ok] template apre: {os.path.relpath(TEMPLATE, ROOT)}")
    except Exception as exc:  # pragma: no cover
        print(f"[FAIL] template non apre: {exc}")
        return 1

    before = corpus_fingerprint()
    os.makedirs(OUT, exist_ok=True)

    rendered = 0
    for name, builder in EXAMPLES.items():
        out_path = os.path.join(OUT, f"example_{name}.docx")
        # Guard: never write into the corpus.
        assert os.path.abspath(out_path).startswith(os.path.abspath(OUT))
        result = render_letter(builder(), TEMPLATE, out_path)
        rendered += 1
        flag = "  needs_review" if result.needs_review else ""
        print(f"[ok] reso example_{name}.docx "
              f"({result.paragraph_count} paragrafi){flag}")
        if result.warnings:
            for w in result.warnings:
                print(f"       - {w}")
        errors.extend(check_output(out_path))

    if rendered < 3:
        errors.append(f"Solo {rendered} strutture rappresentative (servono >= 3)")

    # Corpus untouched.
    after = corpus_fingerprint()
    if before != after:
        changed = [k for k in before if before.get(k) != after.get(k)]
        added = set(after) - set(before)
        errors.append(f"previous_works/ MODIFICATA (changed={changed} added={sorted(added)})")
    else:
        print(f"[ok] previous_works/ intatta ({len(before)} file invariati)")

    print()
    if errors:
        print(f"VALIDAZIONE FALLITA — {len(errors)} problemi:")
        for e in errors:
            print(f"  - {e}")
        return 1
    print(f"VALIDAZIONE OK — {rendered} lettere rese, corpus intatto, "
          f"template ereditato.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
