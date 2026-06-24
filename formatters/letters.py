"""Deterministic composer for Bergamo Legal **letters**.

This module turns *semantic content* (date, recipient, subject, body sections,
closing, signature, attachments) into a formatted ``.docx`` that inherits the
letterhead/header/footer/margins of ``assets/Template_Vuoto.docx``.

It is **content-driven**, not a one-off converter: a developer creates a new
letter by filling :class:`LetterDocument` fields and calling
:func:`render_letter`. No layout logic needs to be rewritten per letter.

The formatting rules implemented here are documented in
``docs/regole_formattazione_lettere.md`` and were inferred from the historical
``letters`` corpus (manifest IDs: 001, 002, 003, 004, 005, 007, 010, 011, 013,
015, 025, 027, 028, 029, 032).

Scope: letters only. Acts/ricorsi/memorie/querele are intentionally unsupported.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from enum import Enum
from typing import Iterable, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

# ---------------------------------------------------------------------------
# Style constants (derived from the corpus — see docs/regole_formattazione_lettere.md)
# ---------------------------------------------------------------------------

BODY_FONT = "Times New Roman"  # HARD_002: explicit on every body run
SIZE_BODY = 12                 # corpo
# Two heading tiers, kept deliberately distinct (client rule):
#  - SIZE_HEADING (16): OGGETTO label + ritual headings (IN FATTO, IN DIRITTO,
#    DICHIARA, DIFFIDA, PREMESSO CHE, ...), always CENTER;
#  - SIZE_SUBHEADING (14): "titoletti" (Fatto, Diritto, MOTIVA, CHIEDE ALTRESÌ,
#    numbered topics), ALWAYS LEFT, with their content on the next line (a capo).
SIZE_HEADING = 16
SIZE_SUBHEADING = 14
SIZE_SMALL = 10                # disclaimer / postscript

RECIPIENT_INDENT_CM = 8.5      # STYLE_008B
LIST_INDENT_CM = 0.5           # default; corpus uses 0.5–1.0

DEFAULT_TEMPLATE = "assets/Template_Vuoto.docx"

# Written in the date position when no date is supplied, so the slot is always
# present and the review path flags it (see collect_warnings).
DATE_PLACEHOLDER = "[INSERISCI QUI LA DATA]"

_AL = WD_ALIGN_PARAGRAPH

# Tokens that mean "a human still has to fill this in".
_PLACEHOLDER_RE = re.compile(r"\[[^\]]*(?:DA INSERIRE|DATA|INSERIRE|\.\.\.|…)[^\]]*\]", re.IGNORECASE)
# Decorative dividers that must never reach the output (STYLE_011).
_DIVIDER_RE = re.compile(r"^\s*(?:\*\s*){2,}\*?\s*$")

# Personal courtesy appellatives that should sit on the same line as the name
# (client feedback 028). "Spett.le" is excluded on purpose: companies keep the
# appellative on its own line (5/5 references 003, 025).
_COURTESY_APPELLATIVES = {
    "egr.", "egr.mo", "egr.ma", "egregio", "egregia",
    "preg.mo", "preg.ma", "pregiatissimo", "pregiatissima",
    "gent.mo", "gent.ma", "gentile", "gentilissimo", "gentilissima",
    "ill.mo", "ill.ma", "illustrissimo", "illustrissima",
    "chiar.mo", "chiar.ma",
}


# ---------------------------------------------------------------------------
# Inline content model
# ---------------------------------------------------------------------------

@dataclass
class Span:
    """A run of text with explicit emphasis. Used for inline bold/italic."""

    text: str
    bold: bool = False
    italic: bool = False


# A "line" of text is either a plain string or a list of Spans (for inline
# emphasis). ``Line`` is the accepted input type everywhere a single styled
# paragraph line is expected.
Line = "str | Span | list[Span]"


class BlockKind(str, Enum):
    """Semantic kinds of body block (see rules doc §4.6)."""

    PARAGRAPH = "paragraph"          # JUSTIFY 12pt
    HEADING_CENTER = "heading_center"  # ritual heading, CENTER 16pt bold
    HEADING_LEFT = "heading_left"      # subheading, LEFT 14pt bold
    LIST_ITEM = "list_item"            # JUSTIFY with left indent


@dataclass
class Block:
    """A single body block. Prefer the helper constructors below."""

    kind: BlockKind = BlockKind.PARAGRAPH
    content: object = ""             # str | Span | list[Span]
    indent_cm: float | None = None   # overrides default for LIST_ITEM
    bold: bool | None = None         # force whole-block emphasis (optional)
    italic: bool | None = None


# Convenience constructors — keep call sites readable.
def paragraph(content) -> Block:
    """A normal justified body paragraph."""
    return Block(BlockKind.PARAGRAPH, content)


def section_heading(content) -> Block:
    """A centered ritual heading (IN FATTO, DIFFIDA, DICHIARA, ...)."""
    return Block(BlockKind.HEADING_CENTER, content)


def subheading(content) -> Block:
    """A left-aligned 14pt bold subheading (Fatto, Diritto, numbered topics)."""
    return Block(BlockKind.HEADING_LEFT, content)


def list_item(content, indent_cm: float | None = None) -> Block:
    """An indented list entry. The marker ((i), a), •, ...) is part of content."""
    return Block(BlockKind.LIST_ITEM, content, indent_cm=indent_cm)


def disposition(verb, content=None) -> list[Block]:
    """A dispositive 'titoletto' rendered per the client rule.

    The ``verb`` becomes a 14 pt **left** subheading on its own line (kept with
    the next paragraph); ``content``, if given, goes **a capo** as a justified
    body paragraph. Use this for inline dispositive verbs that would otherwise
    bundle their continuation on the same line, e.g.::

        disposition("CHIEDE ALTRESÌ", "ai sensi dell'art. 229 CPI ...")

    Returns a list of Blocks; splice it directly into ``body_blocks`` (the
    renderer flattens nested block lists).
    """
    blocks: list[Block] = [Block(BlockKind.HEADING_LEFT, verb)]
    if content is not None and _plain(content).strip():
        blocks.append(Block(BlockKind.PARAGRAPH, content))
    return blocks


# ---------------------------------------------------------------------------
# Letter document model
# ---------------------------------------------------------------------------

class SubjectStyle(str, Enum):
    SPLIT = "split"    # "OGGETTO:" on its own centered line + content line
    INLINE = "inline"  # "OGGETTO: <text>" on a single line


@dataclass
class LetterDocument:
    """Semantic content of one Bergamo Legal letter.

    Fill the fields and pass to :func:`render_letter`. Empty/None optional
    blocks are simply omitted from the output.
    """

    recipient_block: Sequence  # required: list of Line
    subject: object            # required: str | Span | list[Span] (without/with "OGGETTO:")
    opening: object            # required: Line
    body_blocks: Sequence      # required: list of Block | str (str -> paragraph)
    signature_block: Sequence  # required: list of Line
    closing: object = None     # Line | None
    date_place: object = None  # Line | None
    delivery_method: object = None  # Line | None
    attachments: Sequence = field(default_factory=list)  # list of Line
    postscript: object = None  # Line | None (small italic disclaimer)

    # --- variation points (see rules doc §7) ---
    date_above_recipient: bool = True
    delivery_inline_with_recipient: bool = False  # render as first recipient line
    # Default INLINE + justified per client feedback on 027 ("oggetto: testo
    # giustificato") confirmed by the 5/5 reference 025; SPLIT stays available.
    subject_style: SubjectStyle = SubjectStyle.INLINE
    subject_label_center: bool = False  # center the "OGGETTO:" label (split style)
    subject_content_center: bool = False  # center the subject content
    opening_formal: bool = True   # True -> bold + JUSTIFY; False -> plain + LEFT
    list_indent_cm: float = LIST_INDENT_CM
    # Merge a personal courtesy appellative (Egr., Preg.ma, ...) onto the same
    # line as the following name, per client feedback on 028 ("mono linea").
    # "Spett.le" is intentionally excluded (companies stay on their own line,
    # as in the 5/5 references 003 and 025).
    merge_courtesy_appellative: bool = True


@dataclass
class RenderResult:
    """Outcome of a render: where it was written and any review warnings."""

    output_path: str
    warnings: list[str] = field(default_factory=list)
    paragraph_count: int = 0

    @property
    def needs_review(self) -> bool:
        return bool(self.warnings)


# ---------------------------------------------------------------------------
# Text normalization (rules doc §5)
# ---------------------------------------------------------------------------

def normalize_text(text: str) -> str:
    """Apply HARD_004B dash normalization. Placeholders are left untouched."""
    if text is None:
        return ""
    return text.replace("–", "-").replace("—", "-")


def _as_spans(line) -> list[Span]:
    """Coerce a Line (str | Span | list[Span]) into a list of Spans."""
    if line is None:
        return [Span("")]
    if isinstance(line, Span):
        return [line]
    if isinstance(line, str):
        return [Span(line)]
    if isinstance(line, (list, tuple)):
        out: list[Span] = []
        for item in line:
            out.extend(_as_spans(item))
        return out or [Span("")]
    # Fallback: stringify unknown types.
    return [Span(str(line))]


def _plain(line) -> str:
    """Flatten a Line to its plain text (for validation scans)."""
    return "".join(s.text for s in _as_spans(line))


# ---------------------------------------------------------------------------
# Low-level docx helpers
# ---------------------------------------------------------------------------

def _clear_body(document: Document) -> None:
    """Remove all body paragraphs while preserving the section properties
    (margins/page size) and the header/footer parts."""
    for para in list(document.paragraphs):
        para._element.getparent().remove(para._element)


def _add_paragraph(
    document: Document,
    line,
    *,
    align,
    size: int = SIZE_BODY,
    bold: bool | None = None,
    italic: bool | None = None,
    left_indent_cm: float | None = None,
    space_before: float | None = None,
    space_after: float | None = None,
    keep_with_next: bool = False,
):
    """Append one styled paragraph built from a Line.

    ``bold``/``italic`` (when not None) force emphasis on every run; otherwise
    each Span's own emphasis is honoured. Font and size are set explicitly on
    every run (HARD_002 / HARD_003).
    """
    para = document.add_paragraph()
    para.alignment = align
    pf = para.paragraph_format
    if left_indent_cm is not None:
        pf.left_indent = Cm(left_indent_cm)
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if keep_with_next:
        pf.keep_with_next = True

    for span in _as_spans(line):
        text = normalize_text(span.text)
        run = para.add_run(text)
        run.font.name = BODY_FONT
        # Ensure east-asian/complex-script font also maps to TNR.
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(attr), BODY_FONT)
        run.font.size = Pt(size)
        run.bold = span.bold if bold is None else bold
        run.italic = span.italic if italic is None else italic
    return para


# ---------------------------------------------------------------------------
# Block-level composition
# ---------------------------------------------------------------------------

def _coerce_block(item) -> Block:
    if isinstance(item, Block):
        return item
    # Bare string/Span/list -> normal paragraph.
    return Block(BlockKind.PARAGRAPH, item)


def _flatten_blocks(body_blocks) -> list[Block]:
    """Expand nested block lists (e.g. produced by :func:`disposition`) into a
    flat list of Blocks, coercing bare strings/Spans to paragraphs."""
    out: list[Block] = []
    for raw in body_blocks:
        if isinstance(raw, list) and all(isinstance(x, Block) for x in raw):
            out.extend(raw)
        else:
            out.append(_coerce_block(raw))
    return out


def _emit_delivery(document: Document, letter: LetterDocument) -> None:
    _add_paragraph(
        document, letter.delivery_method,
        align=_AL.LEFT, size=SIZE_BODY, italic=True, space_after=8,
    )


def _emit_date(document: Document, line) -> None:
    _add_paragraph(
        document, line,
        align=_AL.RIGHT, size=SIZE_BODY, space_before=6, space_after=10,
    )


def _merge_appellatives(lines: list) -> list:
    """Join a standalone personal courtesy appellative line with the next line
    (client feedback 028: "Egr." + name should be a single line)."""
    merged: list = []
    skip_next = False
    for i, line in enumerate(lines):
        if skip_next:
            skip_next = False
            continue
        word = _plain(line).strip().lower()
        if word in _COURTESY_APPELLATIVES and i + 1 < len(lines):
            head = _as_spans(line)
            if head and not head[-1].text.endswith(" "):
                head = head[:-1] + [Span(head[-1].text + " ", head[-1].bold, head[-1].italic)]
            merged.append([*head, *_as_spans(lines[i + 1])])
            skip_next = True
        else:
            merged.append(line)
    return merged


def _emit_recipient(document: Document, letter: LetterDocument) -> None:
    lines = list(letter.recipient_block)
    if letter.delivery_method is not None and letter.delivery_inline_with_recipient:
        lines = [letter.delivery_method, *lines]
    if letter.merge_courtesy_appellative:
        lines = _merge_appellatives(lines)
    for i, line in enumerate(lines):
        last = i == len(lines) - 1
        _add_paragraph(
            document, line,
            align=_AL.LEFT, size=SIZE_BODY,
            left_indent_cm=RECIPIENT_INDENT_CM,
            space_after=8 if last else 2,
        )


def _emit_subject(document: Document, letter: LetterDocument) -> None:
    spans = _as_spans(letter.subject)
    raw = _plain(spans)
    has_label = raw.strip().upper().startswith("OGGETTO")

    if letter.subject_style is SubjectStyle.INLINE:
        # Single paragraph: label (16pt bold) + content (12pt bold).
        para = document.add_paragraph()
        para.alignment = _AL.CENTER if letter.subject_content_center else _AL.JUSTIFY
        pf = para.paragraph_format
        pf.space_before = Pt(8)
        pf.space_after = Pt(8)
        if has_label:
            label, _, rest = raw.partition(":")
            _add_run(para, normalize_text(label + ":"), size=SIZE_HEADING, bold=True)
            if rest:
                _add_run(para, normalize_text(rest), size=SIZE_BODY, bold=True)
        else:
            _add_run(para, "OGGETTO: ", size=SIZE_HEADING, bold=True)
            for s in spans:
                _add_run(para, normalize_text(s.text), size=SIZE_BODY,
                         bold=True, italic=s.italic)
        return

    # SPLIT style: "OGGETTO:" label line + content line.
    _add_paragraph(
        document, "OGGETTO:",
        align=_AL.CENTER if letter.subject_label_center else _AL.JUSTIFY,
        size=SIZE_HEADING, bold=True, space_before=10, space_after=4,
        keep_with_next=True,
    )
    if has_label:
        # strip the leading "OGGETTO:" from the provided content
        idx = raw.find(":")
        content_spans = _strip_label(spans, idx + 1)
    else:
        content_spans = spans
    _add_paragraph(
        document, content_spans,
        align=_AL.CENTER if letter.subject_content_center else _AL.JUSTIFY,
        size=SIZE_BODY, bold=True, space_after=8,
    )


def _strip_label(spans: list[Span], cut: int) -> list[Span]:
    """Return spans with the first ``cut`` plain-text characters removed."""
    out: list[Span] = []
    remaining = cut
    for s in spans:
        if remaining <= 0:
            out.append(s)
            continue
        if len(s.text) <= remaining:
            remaining -= len(s.text)
            continue
        out.append(Span(s.text[remaining:].lstrip(), s.bold, s.italic))
        remaining = 0
    return out or [Span("")]


def _add_run(para, text, *, size, bold=False, italic=False):
    run = para.add_run(text)
    run.font.name = BODY_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), BODY_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run


def _emit_opening(document: Document, letter: LetterDocument) -> None:
    _add_paragraph(
        document, letter.opening,
        align=_AL.JUSTIFY if letter.opening_formal else _AL.LEFT,
        size=SIZE_BODY,
        bold=True if letter.opening_formal else None,
        space_before=8, space_after=6,
    )


def _emit_body(document: Document, letter: LetterDocument) -> None:
    for block in _flatten_blocks(letter.body_blocks):
        text = _plain(block.content)
        if _DIVIDER_RE.match(text):
            continue  # STYLE_011: drop decorative dividers
        if block.kind is BlockKind.HEADING_CENTER:
            _add_paragraph(
                document, block.content, align=_AL.CENTER,
                size=SIZE_HEADING, bold=True,
                space_before=14, space_after=8, keep_with_next=True,
            )
        elif block.kind is BlockKind.HEADING_LEFT:
            _add_paragraph(
                document, block.content, align=_AL.LEFT,
                size=SIZE_SUBHEADING, bold=True,
                space_before=12, space_after=6, keep_with_next=True,
            )
        elif block.kind is BlockKind.LIST_ITEM:
            _add_paragraph(
                document, block.content, align=_AL.JUSTIFY, size=SIZE_BODY,
                bold=block.bold, italic=block.italic,
                left_indent_cm=block.indent_cm if block.indent_cm is not None
                else letter.list_indent_cm,
                space_after=4,
            )
        else:  # PARAGRAPH
            _add_paragraph(
                document, block.content, align=_AL.JUSTIFY, size=SIZE_BODY,
                bold=block.bold, italic=block.italic, space_after=6,
            )


def _emit_closing(document: Document, letter: LetterDocument) -> None:
    _add_paragraph(
        document, letter.closing,
        align=_AL.RIGHT, size=SIZE_BODY, space_before=16, space_after=6,
    )


def _emit_signature(document: Document, letter: LetterDocument) -> None:
    for i, line in enumerate(letter.signature_block):
        text = _plain(line)
        is_rule = bool(text) and set(text.strip()) <= {"_"}
        is_digital = "firma digitale" in text.lower()
        _add_paragraph(
            document, line, align=_AL.RIGHT, size=SIZE_BODY,
            # STYLE_016: signer name lines bold; rule/role lines plain.
            bold=False if (is_rule or is_digital) else None,
            italic=True if is_digital else None,
            space_before=8 if i == 0 else (4 if not is_rule else None),
            space_after=2,
        )


def _emit_attachments(document: Document, letter: LetterDocument) -> None:
    _add_paragraph(
        document, "Allegati:", align=_AL.LEFT, size=SIZE_BODY, bold=True,
        space_before=16, space_after=4,
    )
    for line in letter.attachments:
        _add_paragraph(
            document, line, align=_AL.LEFT, size=SIZE_BODY,
            left_indent_cm=LIST_INDENT_CM, space_after=3,
        )


def _emit_postscript(document: Document, letter: LetterDocument) -> None:
    _add_paragraph(
        document, letter.postscript, align=_AL.RIGHT, size=SIZE_SMALL,
        italic=True, space_before=4,
    )


# ---------------------------------------------------------------------------
# Validation (rules doc §6)
# ---------------------------------------------------------------------------

def collect_warnings(letter: LetterDocument) -> list[str]:
    """Form-level (not legal) checks that set ``needs_review``."""
    warnings: list[str] = []

    if not letter.recipient_block:
        warnings.append("recipient_block vuoto (blocco obbligatorio).")
    if not _plain(letter.subject).strip():
        warnings.append("subject vuoto (blocco obbligatorio).")
    if not letter.signature_block:
        warnings.append("signature_block vuoto (blocco obbligatorio).")
    if letter.date_place is None or not _plain(letter.date_place).strip():
        warnings.append(
            f"Data assente: inserire la data al posto di {DATE_PLACEHOLDER}."
        )

    # Gather all text actually destined to the document.
    texts: list[str] = []
    for line in letter.recipient_block:
        texts.append(_plain(line))
    texts.append(_plain(letter.subject))
    texts.append(_plain(letter.opening))
    for block in _flatten_blocks(letter.body_blocks):
        texts.append(_plain(block.content))
    for line in letter.signature_block:
        texts.append(_plain(line))
    for line in letter.attachments:
        texts.append(_plain(line))
    for opt in (letter.closing, letter.date_place, letter.delivery_method,
                letter.postscript):
        if opt is not None:
            texts.append(_plain(opt))

    joined = "\n".join(texts)
    placeholders = sorted(set(_PLACEHOLDER_RE.findall(joined)))
    if placeholders:
        warnings.append(
            "Placeholder non risolti da completare: " + ", ".join(placeholders)
        )
    # Check AFTER normalization: render normalizes en/em dash, so a residual
    # here means a dash variant normalize_text does not yet cover (defensive).
    if "–" in normalize_text(joined) or "—" in normalize_text(joined):
        warnings.append("Residuo en/em dash dopo la normalizzazione.")
    for t in texts:
        if _DIVIDER_RE.match(t):
            warnings.append("Divisore decorativo '***' presente (verrà rimosso).")
            break
    return warnings


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def render_letter(
    letter: LetterDocument,
    template_path: str = DEFAULT_TEMPLATE,
    output_path: str = "out/letter.docx",
) -> RenderResult:
    """Compose ``letter`` onto the template and save it to ``output_path``.

    Inherits the template's header (letterhead), footer, margins and page
    setup; only the body is recomposed. Returns a :class:`RenderResult` with
    any ``needs_review`` warnings.
    """
    import os

    document = Document(template_path)  # raises if template missing/unreadable
    _clear_body(document)

    # --- header zone: delivery / date / recipient (order is configurable) ---
    if letter.delivery_method is not None and not letter.delivery_inline_with_recipient:
        _emit_delivery(document, letter)

    # The date is ALWAYS emitted; when missing it is filled with a placeholder
    # so the slot exists and the review path (collect_warnings) flags it.
    date_line = letter.date_place
    if date_line is None or not _plain(date_line).strip():
        date_line = DATE_PLACEHOLDER

    if letter.date_above_recipient:
        _emit_date(document, date_line)
    _emit_recipient(document, letter)
    if not letter.date_above_recipient:
        _emit_date(document, date_line)

    # --- subject / opening (opening is optional: formal istanze have none) ---
    _emit_subject(document, letter)
    if _plain(letter.opening).strip():
        _emit_opening(document, letter)

    # --- body ---
    _emit_body(document, letter)

    # --- closing / signature / attachments / postscript ---
    if letter.closing is not None:
        _emit_closing(document, letter)
    _emit_signature(document, letter)
    if letter.attachments:
        _emit_attachments(document, letter)
    if letter.postscript is not None:
        _emit_postscript(document, letter)

    out_dir = os.path.dirname(output_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    document.save(output_path)

    return RenderResult(
        output_path=output_path,
        warnings=collect_warnings(letter),
        paragraph_count=len(document.paragraphs),
    )


__all__ = [
    "Span",
    "Block",
    "BlockKind",
    "SubjectStyle",
    "LetterDocument",
    "RenderResult",
    "paragraph",
    "section_heading",
    "subheading",
    "list_item",
    "disposition",
    "normalize_text",
    "collect_warnings",
    "render_letter",
]
